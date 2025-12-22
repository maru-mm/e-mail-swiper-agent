"""
Processore per estrarre testo da documenti
"""

import os
from typing import Optional, Tuple
from PyPDF2 import PdfReader
from docx import Document


class DocumentProcessor:
    """
    Estrae testo da vari tipi di documenti
    """
    
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt', '.md']
    
    @staticmethod
    def extract_text(file_path: str, file_type: str = None) -> Tuple[str, str]:
        """
        Estrae il testo da un file
        
        Args:
            file_path: Path del file
            file_type: Tipo di file (opzionale, viene dedotto dall'estensione)
        
        Returns:
            Tupla (testo_estratto, tipo_file)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File non trovato: {file_path}")
        
        # Determina il tipo di file dall'estensione
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            return DocumentProcessor._extract_from_pdf(file_path), 'pdf'
        elif ext in ['.docx', '.doc']:
            return DocumentProcessor._extract_from_docx(file_path), 'docx'
        elif ext in ['.txt', '.md']:
            return DocumentProcessor._extract_from_text(file_path), 'text'
        else:
            raise ValueError(f"Tipo di file non supportato: {ext}")
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """
        Estrae testo da un PDF
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Errore nell'estrazione del PDF: {str(e)}")
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """
        Estrae testo da un file DOCX
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Estrai testo dai paragrafi
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Estrai testo dalle tabelle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Errore nell'estrazione del DOCX: {str(e)}")
    
    @staticmethod
    def _extract_from_text(file_path: str) -> str:
        """
        Legge un file di testo
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Prova con encoding diversi
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                raise Exception(f"Errore nella lettura del file di testo: {str(e)}")
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """
        Verifica se il file è supportato
        
        Args:
            filename: Nome del file
        
        Returns:
            True se supportato
        """
        _, ext = os.path.splitext(filename)
        return ext.lower() in DocumentProcessor.SUPPORTED_EXTENSIONS

